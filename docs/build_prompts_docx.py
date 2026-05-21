"""Genera docs/prompts.docx con los prompts y tecnologías empleadas.

Uso:
    python docs/build_prompts_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).parent / "prompts.docx"


def heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 0:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def body(doc, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def prompt_block(doc, n: int, title: str, prompt: str, result: str) -> None:
    heading(doc, f"Prompt {n}. {title}", level=2)
    body(doc, "Prompt enviado:")
    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Pt(18)
    run = quote.add_run(prompt)
    run.italic = True
    body(doc, f"Resultado: {result}")


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Portada
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Polybilling\nFacturación políglota con arquitectura hexagonal\n")
    r.bold = True
    r.font.size = Pt(22)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "Taller Corte 3 · Programación Avanzada\n"
        "Documento de prompts y tecnologías empleadas\n"
        "Juan Carlos Pastuzán Quintero · ITP · 2026"
    )
    doc.add_page_break()

    # Tecnologías
    heading(doc, "1. Tecnologías empleadas", level=1)
    techs = [
        ("Python 3.12",                "Lenguaje base. Tipado opcional vía dataclasses."),
        ("Flask 3.0",                  "Framework web; blueprints por contexto delimitado."),
        ("Jinja2 + Bootstrap 5",       "Capa de presentación HTML."),
        ("python-dotenv",              "Lectura de .env para configuración por entorno."),
        ("Apache Cassandra 5",         "Almacén de personas (clientes y empleados). "
                                       "Modelado denormalizado con tabla person + person_by_role."),
        ("cassandra-driver 3.29",      "Driver oficial DataStax."),
        ("MongoDB 7",                  "Catálogo de productos. Documentos JSON con índices "
                                       "únicos por SKU y secundarios por categoría/tags."),
        ("pymongo 4.8",                "Driver oficial."),
        ("MySQL 8",                    "Encabezado y detalle de factura, transaccional ACID."),
        ("mysql-connector-python 9",   "Driver oficial Oracle, con pool de conexiones."),
        ("Neo4j 5",                    "Grafo Cliente–Producto–Contexto para recomendaciones."),
        ("neo4j-python-driver 5.23",   "Driver oficial Bolt."),
        ("python-docx 1.1",            "Generación de este mismo documento Word."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Tecnología"
    hdr[1].text = "Uso en el proyecto"
    for name, use in techs:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = use

    doc.add_page_break()

    # Prompts
    heading(doc, "2. Prompts empleados durante la construcción", level=1)
    body(
        doc,
        "A continuación se presentan los prompts (en lenguaje natural) que "
        "se enviaron al asistente de IA, junto con un resumen del resultado "
        "obtenido. Los prompts se redactaron de forma incremental, "
        "validando cada bloque antes de continuar con el siguiente."
    )

    prompts = [
        (
            "Definición de arquitectura y alcance",
            "Diseña una aplicación Flask en Python 3.12 con arquitectura "
            "hexagonal (puertos y adaptadores) que cumpla un taller de "
            "facturación políglota con Cassandra (personas), MongoDB "
            "(productos), MySQL (factura/detalle) y Neo4j (recomendaciones). "
            "Lista la estructura de carpetas, qué va en cada capa y cómo "
            "se cablean las dependencias.",
            "Se obtuvo la estructura del proyecto, el wiring vía un Container "
            "y la justificación de qué motor usar para cada agregado."
        ),
        (
            "Entidades del dominio",
            "Crea entidades del dominio con validaciones de invariantes "
            "(rol, estrato 1-6, salario obligatorio para EMPLOYEE, cantidad "
            "positiva, stock no negativo) usando dataclasses. Incluye Person, "
            "Product, Invoice, InvoiceDetail con propiedades calculadas "
            "(subtotal, IVA, total, line_total).",
            "Se generaron entidades con __post_init__ para validar y "
            "métodos como discount_stock() y add_detail()."
        ),
        (
            "Adaptador Cassandra",
            "Implementa un repositorio Cassandra para personas con dos tablas "
            "(person por documento y person_by_role para listar por rol sin "
            "ALLOW FILTERING). Maneja conversión de tipos (date, UUID) y "
            "evita N+1 al buscar por rol.",
            "Adaptador completo con CassandraConnection singleton, save/find/"
            "delete y mapeo bidireccional de filas a entidades."
        ),
        (
            "Adaptador MongoDB",
            "Repositorio Mongo para productos con índices únicos por SKU, "
            "secundarios por category y tags, upsert por save y mapeo desde "
            "documento BSON a entidad Product.",
            "MongoProductRepository con create_index en el constructor, "
            "upsert con $set y conversión de _id a string."
        ),
        (
            "Adaptador MySQL transaccional",
            "Repositorio MySQL que persista factura y detalle_factura en una "
            "sola transacción (commit/rollback) y permita consultar facturas "
            "con sus líneas y los SKUs comprados por un cliente.",
            "MySQLInvoiceRepository con pool de conexiones, executemany para "
            "el detalle y método find_purchased_skus_by_client."
        ),
        (
            "Motor de recomendaciones Neo4j",
            "Diseña el grafo Cliente-Producto-Contexto: nodos Person, "
            "Product, Category, Tag, Neighborhood, Municipality, Department, "
            "Gender, Stratum; relaciones BOUGHT, LIVES_IN, IN_CATEGORY, "
            "HAS_TAG, HAS_GENDER, HAS_STRATUM, PART_OF. Cypher de "
            "recomendación que combine filtrado colaborativo geográfico, "
            "similitud por categoría/tags y popularidad barrial, devolviendo "
            "un score y la razón principal.",
            "Esquema del grafo y consulta Cypher unificada que devuelve sku, "
            "name, price, score y reason ordenado descendente."
        ),
        (
            "Servicios de aplicación",
            "Casos de uso que orquesten los repositorios: PersonService y "
            "ProductService deben sincronizar Cassandra/Mongo con el grafo "
            "(upsert). InvoiceService debe validar cliente y empleado en "
            "Cassandra, descontar stock en Mongo y registrar la compra en "
            "Neo4j tras persistir en MySQL.",
            "Servicios que reciben los puertos por constructor, no conocen "
            "implementaciones concretas y disparan los efectos colaterales "
            "en el orden correcto."
        ),
        (
            "Interfaz Flask",
            "Crea blueprints por contexto (persons, products, invoices, "
            "recommendations), un Container para inyectar servicios y "
            "templates Jinja2 con Bootstrap 5 y mensajes flash. La pantalla "
            "de nueva factura debe permitir añadir/quitar líneas en cliente.",
            "App factory, 4 blueprints, contenedor de DI y templates "
            "responsive con tabla dinámica de líneas en JS vanilla."
        ),
        (
            "Datos semilla",
            "Genera scripts de seed para los 4 motores con datos coherentes: "
            "5 clientes y 2 empleados en Cassandra, 12 productos en Mongo, "
            "4 facturas en MySQL y 9 compras en Neo4j para alimentar el "
            "motor de recomendaciones desde el primer arranque.",
            "Carpeta dbs/ con 01_cassandra.cql, 02_mongo.js, 03_mysql.sql, "
            "04_neo4j.cypher y seed_all.py que ejecuta los cuatro."
        ),
        (
            "Documentación y empaquetado",
            "Redacta un README profesional con tabla de motores, "
            "instrucciones de instalación, comandos de seed por opción A/B, "
            "y diagrama ASCII de la arquitectura. Genera .gitignore para "
            "Python/Flask, .env.example y requirements.txt fijado a "
            "versiones estables.",
            "README.md de 150 líneas, .gitignore completo, .env.example y "
            "requirements.txt con 12 dependencias bloqueadas."
        ),
    ]

    for i, (title, prompt, result) in enumerate(prompts, 1):
        prompt_block(doc, i, title, prompt, result)

    doc.add_page_break()
    heading(doc, "3. Resumen", level=1)
    body(
        doc,
        "El asistente de IA aceleró la generación de boilerplate y permitió "
        "concentrar el esfuerzo humano en las decisiones de diseño "
        "(qué motor para qué agregado, qué heurística para recomendar, "
        "cómo cablear las dependencias). Cada bloque fue revisado y ajustado "
        "manualmente antes de continuar, garantizando que el resultado "
        "respetara la arquitectura hexagonal y las buenas prácticas de cada "
        "motor de persistencia."
    )

    doc.save(OUT)
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    build()
